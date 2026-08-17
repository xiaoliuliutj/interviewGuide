from io import BytesIO
import asyncio

import pytest

from agent.RAG.ragDocumentParser import DocumentParser
from agent.WorkFlows.Interview.interviewModels import (
    InterviewAction,
    InterviewEvaluation,
    InterviewPlan,
    InterviewPlanStage,
    InterviewRoute,
    InterviewSessionState,
    InterviewStage,
    WorkflowIntentDecision,
)
from agent.WorkFlows.Interview.interviewWorkflow import InterviewWorkflow
from agent.WorkFlows.workflowService import WorkflowService
from agent.Common.AgentRequest import AgentOperationRequest
from agent.Memory.memoryModels import LongTermMemorySnapshot
from agent.WorkFlows.Resume.resumeWorkflow import ResumeWorkflow


def buildPlan() -> InterviewPlan:
    """构造符合六阶段硬约束的计划，用于验证工作流政策不会依赖 LLM 随机输出。"""
    stages = []
    for stage in InterviewStage:
        stages.append(
            InterviewPlanStage(
                stage=stage,
                topics=[f"{stage.value}主题"],
                difficulty="MEDIUM",
                maxPrimaryQuestions=1 if stage in {InterviewStage.OPENING, InterviewStage.SUMMARY} else 4,
                maxFollowupsPerQuestion=0 if stage in {InterviewStage.OPENING, InterviewStage.SUMMARY, InterviewStage.CODING} else 2,
                timeBudgetMinutes=5,
            )
        )
    return InterviewPlan(
        candidateSummary="候选人具备后端项目经验。",
        strategySummary="按项目、基础和工程实践逐步考察。",
        stages=stages,
        selectedSkills=["interview-coach"],
        coverageMatrix={
            "project_or_internship": True,
            "technical_stack": True,
            "knowledge_and_practice": True,
        },
    )


def testInterviewPlanRejectsInvalidStageOrder() -> None:
    """验证模型即使返回六个阶段，也不能交换阶段顺序绕过工作流。"""
    plan = buildPlan()
    invalid = plan.model_dump(mode="json")
    invalid["stages"][0], invalid["stages"][1] = invalid["stages"][1], invalid["stages"][0]
    with pytest.raises(ValueError):
        InterviewPlan.model_validate(invalid)


def testCodingStageDoesNotIssueSecondQuestionAfterPassing() -> None:
    """验证参考流程中的算法题限制：首题达到及格线后不能继续增加第二题。"""
    workflow = InterviewWorkflow(None, None, None, None)
    state = InterviewSessionState(
        sessionId="session-1",
        userId="user-1",
        targetRole="Java后端",
        plan=buildPlan(),
        currentStage=InterviewStage.CODING,
        currentTopic="数组",
        currentQuestion="请说明二分查找。",
        primaryQuestionCount=1,
        totalPrimaryQuestionCount=8,
        totalQuestionCount=8,
        stageQuestionCounts={InterviewStage.CODING.value: 1},
    )
    evaluation = InterviewEvaluation(
        evaluationSummary="回答正确。",
        score=80,
        answerSummary="说明了核心思路。",
    )
    allowed = workflow.getAllowedActions(state, evaluation)
    assert InterviewAction.NEXT_QUESTION not in allowed


def testTransitionIntoSummaryBecomesEndInterview() -> None:
    """验证 SUMMARY 不会错误生成新的候选人作答题，而是直接进入最终总结。"""
    workflow = InterviewWorkflow(None, None, None, None)
    state = InterviewSessionState(
        sessionId="session-1",
        userId="user-1",
        targetRole="Java后端",
        plan=buildPlan(),
        currentStage=InterviewStage.CODING,
        currentQuestion="算法。",
        stageQuestionCounts={InterviewStage.CODING.value: 2},
    )
    evaluation = InterviewEvaluation(
        evaluationSummary="需要总结。",
        score=30,
        answerSummary="回答不完整。",
    )
    route = workflow.normalizeRoute(
        state,
        evaluation,
        InterviewRoute(action=InterviewAction.NEXT_STAGE, nextTopic=None),
        {InterviewAction.NEXT_STAGE},
    )
    assert route.action == InterviewAction.END_INTERVIEW


def testDocxParserExtractsHeadingAndText() -> None:
    """验证简。"DOCX 会被真实解析为文本，而不是作为未实现格式被跳过。"""
    from docx import Document

    document = Document()
    document.add_heading("工作经历", level=1)
    document.add_paragraph("负责 Java 服务端开发。")
    buffer = BytesIO()
    document.save(buffer)
    sections = DocumentParser().parseDocx(buffer.getvalue())
    assert sections[0].headingPath == "工作经历"
    assert "Java" in sections[0].content


def testNaturalLanguageRouterUsesRestrictedWorkflowSchema() -> None:
    """验证顶层路由只能返回注册工作流，且会解析模型 JSON 而非依赖 Java 任务码。"""
    class FakeLlm:
        async def requestJson(self, messages, temperature):
            return {"workflow": "INTERVIEW", "intent": "START_INTERVIEW", "confidence": 0.9}

    runtime = WorkflowService(FakeLlm(), None, None, None, None, None)
    request = AgentOperationRequest.model_validate({
        "context": {
            "apiVersion": "v1",
            "requestId": "request-1",
            "runId": "run-1",
            "principalId": "user-1",
            "conversationId": "conversation-1",
            "timestamp": "2026-08-17T10:00:00Z",
        },
            "prompt": "???? Java ???????",
    })
    decision = asyncio.run(runtime.resolveWorkflow(request))
    assert decision.workflow == "INTERVIEW"


def testInterviewInitializationPersistsOnlyAfterPlanValidation() -> None:
    """验证自然语言启动面试会先生成合规计划，再提交开场状态和统一响应。"""
    class FakeLlm:
        async def requestJson(self, messages, temperature):
            return buildPlan().model_dump(mode="json")

    class FakeMemoryRepository:
        async def loadLongTermMemory(self, userId, resumeId):
            return LongTermMemorySnapshot(None, None, None)

        async def loadResumeMemoryDetail(self, userId, resumeId):
            return None

    class FakeMemoryService:
        repository = FakeMemoryRepository()

    class FakeRepository:
        def __init__(self):
            self.committed = None

        async def loadRunResult(self, runId, sessionId, userId):
            return None

        async def loadState(self, sessionId, userId):
            return None

        async def ensureSession(self, sessionId, userId, resumeId):
            return None

        async def claimRun(self, sessionId, userId, runId, version):
            return "PROCESSING"

        async def commitState(self, state, runId, expectedVersion, responseJson, **values):
            self.committed = state

        async def failRun(self, sessionId, runId, responseJson):
            raise AssertionError("初始化不应进入失败路。")

    repository = FakeRepository()
    workflow = InterviewWorkflow(FakeLlm(), FakeMemoryService(), None, repository)
    request = AgentOperationRequest.model_validate({
        "context": {
            "apiVersion": "v1",
            "requestId": "request-1",
            "runId": "run-1",
            "principalId": "user-1",
            "conversationId": "conversation-1",
            "timestamp": "2026-08-17T10:00:00Z",
        },
        "prompt": "开始模拟面试。",
        "data": {"targetRole": "Java后端", "resumeId": "resume-1"},
    })
    response = asyncio.run(workflow.handleRequest(request))
    assert response.data["type"] == "INTERVIEW_QUESTION"
    assert repository.committed.currentStage == InterviewStage.OPENING
    assert repository.committed.stateVersion == 1


def testInterviewAnswerFollowsEvaluationRouteThenQuestionGeneration() -> None:
    """验证一轮面试严格按评价、路由、检索、出题和提交的顺序推进。"""
    class FakeLlm:
        def __init__(self):
            self.responses = [
                {
                    "evaluationSummary": "回答说明了线程池核心参数。",
                    "score": 70,
                    "answerSummary": "说明了核心线程数和队列。",
                    "strengths": ["基础概念正确"],
                    "weaknesses": ["缺少拒绝策略细节"],
                },
                {"action": "NEXT_QUESTION", "nextTopic": "线程池拒绝策。"},
                {"question": "请说明线程池拒绝策略的适用场景。"},
            ]

        async def requestJson(self, messages, temperature):
            return self.responses.pop(0)

    class FakeMemoryRepository:
        async def loadLongTermMemory(self, userId, resumeId):
            return LongTermMemorySnapshot(None, None, None)

        async def loadResumeMemoryDetail(self, userId, resumeId):
            return None

    class FakeMemoryService:
        repository = FakeMemoryRepository()

    class FakeRepository:
        def __init__(self):
            self.turn = None

        async def claimRun(self, sessionId, userId, runId, version):
            return "PROCESSING"

        async def loadRunResult(self, runId, sessionId, userId):
            return None

        async def commitState(self, state, runId, expectedVersion, responseJson, *values):
            self.turn = state.turns[-1]

        async def failRun(self, sessionId, runId, responseJson):
            raise AssertionError("回答路径不应失败")

    state = InterviewSessionState(
        sessionId="conversation-1",
        userId="user-1",
        targetRole="Java后端",
        plan=buildPlan(),
        currentStage=InterviewStage.FUNDAMENTAL,
        currentTopic="线程。",
        currentQuestion="请说明线程池核心参数。",
        stateVersion=1,
        primaryQuestionCount=1,
        totalPrimaryQuestionCount=2,
        totalQuestionCount=2,
        stageQuestionCounts={InterviewStage.FUNDAMENTAL.value: 1},
    )
    repository = FakeRepository()
    workflow = InterviewWorkflow(FakeLlm(), FakeMemoryService(), None, repository)
    request = AgentOperationRequest.model_validate({
        "context": {
            "apiVersion": "v1",
            "requestId": "request-2",
            "runId": "run-2",
            "principalId": "user-1",
            "conversationId": "conversation-1",
            "timestamp": "2026-08-17T10:00:00Z",
        },
        "prompt": "核心线程数决定常驻线程数量，队列用于暂存任务。",
        "stateVersion": 1,
    })
    response = asyncio.run(workflow.handleAnswer(request, state))
    assert response.data["content"] == "请说明线程池拒绝策略的适用场景。"
    assert repository.turn.action == InterviewAction.NEXT_QUESTION
    assert state.stateVersion == 2


def testResumeWorkerParsesThenEvaluatesAndPersistsMemory() -> None:
    """验证简历异。"worker 按“文件解析→LLM评估→长期记忆”顺序执行真实闭环。"""
    class FakeLlm:
        async def requestJson(self, messages, temperature):
            return {
                "overallScore": 80,
                "contentScore": 81,
                "structureScore": 79,
                "skillMatchScore": 82,
                "expressionScore": 78,
                "projectScore": 80,
                "summary": "简历与目标岗位基本匹配。",
                "strengths": ["项目经历清晰"],
                "suggestions": ["补充量化结果"],
                "issues": [{"question": "请补充职责边。", "priority": "MEDIUM", "suggestion": "明确个人贡献"}],
                "technicalStack": ["Java"],
                "technicalDepth": ["Spring"],
                "careerPreferences": ["后端开。"],
            }

    class FakeRepository:
        def __init__(self):
            self.completed = None

        async def claimJobs(self):
            return [{
                "run_id": "resume-run-1",
                "resume_id": "resume-1",
                "user_id": "user-1",
                "target_role": "Java后端",
                "attempt_count": 0,
            }]

        async def loadDocument(self, resumeId, userId):
            return {
                "raw_content": b"# \xe6\x8a\x80\xe6\x9c\xaf\xe6\xa0\x88\nJava \xe5\x90\x8e\xe7\xab\xaf\xe9\xa1\xb9\xe7\x9b\xae\xe5\xbc\x80\xe5\x8f\x91",
                "file_name": "resume.md",
                "content_type": "text/markdown",
            }

        async def markAnalyzing(self, resumeId, text):
            assert "Java" in text

        async def completeJob(self, runId, resumeId, evaluation):
            self.completed = evaluation

        async def failJob(self, *values):
            raise AssertionError("有效简历不应进入失败路。")

    class FakeMemory:
        async def saveResumeEvaluation(self, userId, resumeId, evaluation):
            assert evaluation["overallScore"] == 80

    repository = FakeRepository()
    workflow = ResumeWorkflow(FakeLlm(), FakeMemory(), repository)
    asyncio.run(workflow.processJobs())
    assert repository.completed.overallScore == 80
