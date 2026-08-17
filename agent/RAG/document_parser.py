import re
from dataclasses import dataclass
from io import BytesIO

from agent.Common.Exceptions.agent_exception import RagDocumentParseError, RagDocumentTooLargeError
from agent.RAG.policy import RagPolicy


@dataclass(frozen=True)
class ParsedSection:
    content: str
    headingPath: str | None = None
    pageNumber: int | None = None


class DocumentParser:
    """将上传的原始文件转换为带标题和页码元数据的纯文本段落。"""

    def __init__(self, policy: RagPolicy | None = None) -> None:
        self.policy = policy or RagPolicy()

    def parse(self, content: bytes, fileName: str, contentType: str | None = None) -> list[ParsedSection]:
        """根据扩展名解析 Markdown、PDF 或纯文本；解析失败会抛出明确的领域异常。"""
        if len(content) > self.policy.maxDocumentBytes:
            raise RagDocumentTooLargeError("文档超过允许的大小限制")
        suffix = fileName.lower().rsplit(".", 1)[-1] if "." in fileName else "txt"
        try:
            if suffix in {"md", "markdown"} or contentType == "text/markdown":
                return self.parseMarkdown(content.decode("utf-8"))
            if suffix == "pdf" or contentType == "application/pdf":
                return self.parsePdf(content)
            if suffix == "docx" or contentType == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self.parseDocx(content)
            return [ParsedSection(content.decode("utf-8"))]
        except RagDocumentParseError:
            raise
        except Exception as error:
            raise RagDocumentParseError(f"无法解析文档 {fileName}") from error

    def parseMarkdown(self, text: str) -> list[ParsedSection]:
        """按 Markdown 标题建立 section，并保留完整标题路径供追踪。"""
        sections: list[ParsedSection] = []
        headings: list[str] = []
        body: list[str] = []
        def flush() -> None:
            value = "\n".join(body).strip()
            if value:
                sections.append(ParsedSection(value, " / ".join(headings) or None))
            body.clear()
        for line in text.splitlines():
            match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
            if match:
                flush()
                level = len(match.group(1))
                headings[:] = headings[: level - 1]
                headings.append(match.group(2))
            else:
                body.append(line)
        flush()
        return sections or [ParsedSection(text.strip())]

    def parsePdf(self, content: bytes) -> list[ParsedSection]:
        """使用 pypdf 提取逐页文本；无法提取时明确报告解析错误。"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            sections: list[ParsedSection] = []
            for index, page in enumerate(reader.pages):
                pageText = page.extract_text() or ""
                sections.extend(self.splitPdfPage(pageText, index + 1))
            if not sections:
                raise RagDocumentParseError("PDF 未提取到文本")
            return sections
        except ImportError as error:
            raise RagDocumentParseError("缺少 pypdf 依赖") from error

    def parseDocx(self, content: bytes) -> list[ParsedSection]:
        """提取 DOCX 段落和标题，供简历解析及知识库索引共用同一份可靠文本。"""
        try:
            from docx import Document

            document = Document(BytesIO(content))
            sections: list[ParsedSection] = []
            headingPath: list[str] = []
            body: list[str] = []

            def flush() -> None:
                """在遇到新标题或文档结束时提交已有正文，避免标题和正文丢失关联。"""
                value = "\n".join(body).strip()
                if value:
                    sections.append(ParsedSection(value, " / ".join(headingPath) or None))
                body.clear()

            for paragraph in document.paragraphs:
                value = paragraph.text.strip()
                if not value:
                    continue
                styleName = getattr(paragraph.style, "name", "") or ""
                if styleName.lower().startswith("heading"):
                    flush()
                    levelText = styleName.split()[-1]
                    level = int(levelText) if levelText.isdigit() else 1
                    headingPath[:] = headingPath[: max(level - 1, 0)]
                    headingPath.append(value)
                else:
                    body.append(value)
            flush()
            if not sections:
                text = "\n".join(item.text for item in document.paragraphs).strip()
                if not text:
                    raise RagDocumentParseError("DOCX 未提取到有效文本")
                return [ParsedSection(text)]
            return sections
        except ImportError as error:
            raise RagDocumentParseError("缺少 python-docx 依赖") from error

    def splitPdfPage(self, pageText: str, pageNumber: int) -> list[ParsedSection]:
        """识别 PDF 中常见的章节编号或 Markdown 标题，识别不到时按页返回。"""
        lines = pageText.splitlines()
        heading = None
        body: list[str] = []
        sections: list[ParsedSection] = []
        for line in lines:
            value = line.strip()
            if re.match(r"^(第[一二三四五六七八九十百]+章|\d+(?:\.\d+)*\s+|#{1,6}\s+)", value):
                if body and "\n".join(body).strip():
                    sections.append(ParsedSection("\n".join(body).strip(), heading, pageNumber))
                    body.clear()
                heading = value.lstrip("# ")
            else:
                body.append(line)
        if body and "\n".join(body).strip():
            sections.append(ParsedSection("\n".join(body).strip(), heading, pageNumber))
        return sections or [ParsedSection(pageText, None, pageNumber)]
